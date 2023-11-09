import datetime
import os
import pdfrw
import io
import locale
import requests

from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from datetime import date
from PyPDF2 import PdfFileReader, PdfFileWriter
from pdf2image import convert_from_path
from PIL import Image

from django.utils import timezone

from applications.common.models import SiteSettings
from applications.core.models import ContractAdmin, Tariff


class DocumentGenerator:
    ANNOT_KEY = '/Annots'
    ANNOT_FIELD_KEY = '/T'
    ANNOT_VAL_KEY = '/V'
    ANNOT_RECT_KEY = '/Rect'
    SUBTYPE_KEY = '/Subtype'
    WIDGET_SUBTYPE_KEY = '/Widget'

    def fill_pdf(self, input_pdf_path, output_pdf_path, data_dict):
        template_pdf = pdfrw.PdfReader(input_pdf_path)
        for page in template_pdf.pages:
            annotations = page[self.ANNOT_KEY]
            for annotation in annotations:
                if annotation[self.SUBTYPE_KEY] == self.WIDGET_SUBTYPE_KEY:
                    if annotation[self.ANNOT_FIELD_KEY]:
                        key = annotation[self.ANNOT_FIELD_KEY][1:-1]
                        if key in data_dict.keys():
                            if type(data_dict[key]) == bool:
                                if data_dict[key]:
                                    annotation.update(pdfrw.PdfDict(
                                        AS=pdfrw.PdfName('1')))
                            else:
                                annotation.update(
                                    pdfrw.PdfDict(V='{}'.format(data_dict[key]))
                                )
                                annotation.update(pdfrw.PdfDict(AP=''))
        pdfrw.PdfWriter().write(output_pdf_path, template_pdf)

    def __init__(self, profile):
        self.profile = profile

        dt = timezone.localtime().strftime('%d.%m.%Y')
        filename_pdf = f'{self.profile.last_name}-{self.profile.first_name}-immatrikulation-{dt}.pdf'
        filename_jpg = f'{self.profile.last_name}-{self.profile.first_name}-immatrikulation-{dt}.jpg'

        self.BASE_PATH = os.path.dirname(os.path.dirname(os.path.abspath(__file__).replace('/applications', '')))
        self.COMMON_PATH = os.path.join(self.BASE_PATH, 'common_docs/')
        self.MEDIA_PATH = os.path.join(self.BASE_PATH, 'media/')

        self.TEMPLATE_PATH = f'{self.COMMON_PATH}documents/immat.pdf'
        self.OUTPUT_PATH = f'{self.MEDIA_PATH}documents/starter_form.pdf'
        self.FINAL_PDF_PATH = f'{self.MEDIA_PATH}documents/{filename_pdf}'
        self.FINAL_JPG_PATH = f'{self.MEDIA_PATH}documents/{filename_jpg}'

    def fill_starter_checklist(self):
        self.today = date.today()
        data_dict = {
            'FEFF00420065007A0065006900630068006E0075006E0067005B0030005D': self.profile.university.name_de,
            'FEFF0041006E0073006300680072006900660074005B0030005D': self.profile.university.address,
            'FEFF00540065006C00650066006F006E006E0075006D006D00650072005B0030005D': f'{self.profile.university.phone}, {self.profile.university.site}',
            'FEFF004E0061006D00650056006F0072006E0061006D0065005B0030005D': f'{self.profile.last_name} {self.profile.first_name}',
            'FEFF00470065006200750072007400730064006100740075006D002D005400610067005B0030005D': self.profile.bday.strftime("%d"),
            'FEFF00470065006200750072007400730064006100740075006D002D004D006F006E00610074005B0030005D': self.profile.bday.strftime("%m"),
            'FEFF00470065006200750072007400730064006100740075006D002D004A006100680072005B0030005D': self.profile.bday.strftime("%Y"),
            'FEFF0053007400610061007400730061006E00670065006800F6007200690067006B006500690074005B0030005D': self.profile.nationality,
            'FEFF005300740075006400690065006E0066006100630068005B0030005D': self.profile.faculty.name_de,
            'FEFF00650069006E0067006500730063006800720069006500620065006E002D005400610067005B0030005D': '01', # profile.study_start.day,
            'FEFF00650069006E0067006500730063006800720069006500620065006E002D004D006F006E00610074005B0030005D': '09', #profile.study_start.month,
            'FEFF00650069006E0067006500730063006800720069006500620065006E002D004A006100680072005B0030005D': self.profile.study_start.strftime("%y"),
            'FEFF0052006500670065006C00640061007500650072005B0030005D': self.profile.study_end.year-self.profile.study_start.year,
            'FEFF005300740075006400690065006E0065006E00640065002D005400610067005B0030005D': '30', # profile.study_end.day,
            'FEFF005300740075006400690065006E0065006E00640065002D004D006F006E00610074005B0030005D': '06', # profile.study_end.month,
            'FEFF00460065007200690065006E006200650073006300680061006500660074006900670075006E0067005B0030005D': True, # employment_during_vacation
            'FEFF00440061007500650072005400610067005B0030005D': self.profile.summer_holidays_start.strftime('%d'), # '01', # summer_winter_vacation_begin_DAY
            'FEFF00440061007500650072004D006F006E00610074005B0030005D': self.profile.summer_holidays_start.strftime('%m'), #'06', #'summer_winter_vacation_begin_MONTH
            'FEFF00440061007500650072002D006200690073002D005400610067005B0030005D': self.profile.summer_holidays_end.strftime('%d'), # '31', # summer_winter_vacation_end_DAY
            'FEFF00440061007500650072002D006200690073002D004D006F006E00610074005B0030005D': self.profile.summer_holidays_end.strftime('%m'), # '08', # summer_winter_vacation_end_MONTH
            'FEFF006A0061002D006E00650069006E005B0030005D': self.profile.study_end.year > self.today.year, # study_continued_after_vacation-yes
            'FEFF006A0061002D006E00650069006E005B0031005D': self.profile.study_end.year <= self.today.year, # study_continued_after_vacation-no
            'FEFF0044006100740075006D005B0030005D': self.today.strftime("%d.%m.%Y"),
        }

        return self.fill_pdf(self.TEMPLATE_PATH, self.OUTPUT_PATH, data_dict)

    def run(self):
        canvas_data = self.get_overlay_canvas()
        form = self.merge(canvas_data, template_path=self.OUTPUT_PATH)
        self.save(form, filename=self.FINAL_PDF_PATH)

    def rotate_pages(self, pdf_path):
        pdf_writer = PdfFileWriter()
        pdf_reader = PdfFileReader(pdf_path)
        # Rotate page 90 degrees to the right
        page_1 = pdf_reader.getPage(0).rotateClockwise(90)
        pdf_writer.addPage(page_1)

    def get_overlay_canvas(self) -> io.BytesIO:
        data = io.BytesIO()
        pdf = canvas.Canvas(data, pagesize=(5080 * mm, 5080 * mm))
        pdf.setFont('Helvetica', 10)
        pdf.drawString(x=291, y=243, text=self.profile.summer_holidays_start.strftime("%y")) # summer winter vacation start year
        pdf.drawString(x=382, y=243, text=self.profile.summer_holidays_end.strftime("%y")) # summer winter vacation end year
        pdf.drawString(x=777, y=329, text=self.profile.study_end.strftime("%y")) # study end year
        pdf.save()
        data.seek(0)
        return data

    def merge(self, overlay_canvas: io.BytesIO, template_path: str) -> io.BytesIO:
        template_pdf = pdfrw.PdfReader(template_path)
        overlay_pdf = pdfrw.PdfReader(overlay_canvas)
        for page, data in zip(template_pdf.pages, overlay_pdf.pages):
            overlay = pdfrw.PageMerge().add(data)[0]
            pdfrw.PageMerge(page).add(overlay).render()
        form = io.BytesIO()
        pdfrw.PdfWriter().write(form, template_pdf)
        form.seek(0)
        return form

    def save(self, form: io.BytesIO, filename: str):
        with open(filename, 'wb') as f:
            f.write(form.read())

    def generate_document(self):
        if self.profile.is_confirmed and self.profile.is_admin_confirmed and self.profile.is_form_completed\
                and self.profile.summer_holidays_start and self.profile.summer_holidays_end:
            self.fill_starter_checklist()
            self.run()
            pages = convert_from_path(self.FINAL_PDF_PATH, 500)
            for page in pages:
                page.save(self.FINAL_JPG_PATH, 'JPEG')
            os.remove(self.OUTPUT_PATH)
            os.remove(self.FINAL_PDF_PATH)
            return self.FINAL_JPG_PATH
        return None

    def generate_pdf(self):
        if self.profile.is_confirmed and self.profile.is_admin_confirmed and self.profile.is_form_completed\
                and self.profile.summer_holidays_start and self.profile.summer_holidays_end:
            self.fill_starter_checklist()
            self.run()
            os.remove(self.OUTPUT_PATH)
            return self.FINAL_PDF_PATH
        return None


class BewerbungGenerator:
    def __init__(self, profile):
        self.profile = profile

        self.today = timezone.localtime()
        dt = self.today.strftime('%d.%m.%Y-%H:%M:%S')
        filename_pdf = f'{self.profile.last_name}-{self.profile.first_name}-{self.profile.user.email}-jeople-{dt}.pdf'

        self.BASE_PATH = os.path.dirname(os.path.dirname(os.path.abspath(__file__).replace('/applications', '')))
        self.COMMON_PATH = os.path.join(self.BASE_PATH, 'common_docs/')
        self.MEDIA_PATH = os.path.join(self.BASE_PATH, 'media/')

        self.TEMPLATE_PATH1 = f'{self.COMMON_PATH}documents/document-page0.pdf'
        self.TEMPLATE_PATH2 = f'{self.COMMON_PATH}documents/document-page1.pdf'
        self.TEMPLATE_PATH3 = f'{self.COMMON_PATH}documents/document-page2.pdf'

        self.IMG_PATH = f'{self.COMMON_PATH}documents/DS.jpg'

        self.OUTPUT_PATH1 = f'{self.MEDIA_PATH}documents/doc_page0.pdf'
        self.OUTPUT_PATH2 = f'{self.MEDIA_PATH}documents/doc_page1.pdf'
        self.OUTPUT_PATH3 = f'{self.MEDIA_PATH}documents/doc_page2.pdf'

        self.FINAL_PDF_PATH = f'{self.MEDIA_PATH}documents/{filename_pdf}'

        year = self.today.year
        self.document_year = str(year) if self.today.date() < datetime.datetime.strptime(f"01/09/{year}", "%d/%m/%Y").date() else str(year + 1)

    def run(self, canvas_page, template_path, filename):
        canvas_data = self.get_overlay_canvas(canvas_page)
        form = self.merge(canvas_data, template_path=template_path)
        self.save(form, filename=filename)

    def merge(self, overlay_canvas: io.BytesIO, template_path: str) -> io.BytesIO:
        template_pdf = pdfrw.PdfReader(template_path)
        overlay_pdf = pdfrw.PdfReader(overlay_canvas)
        for page, data in zip(template_pdf.pages, overlay_pdf.pages):
            overlay = pdfrw.PageMerge().add(data)[0]
            pdfrw.PageMerge(page).add(overlay).render()
        form = io.BytesIO()
        pdfrw.PdfWriter().write(form, template_pdf)
        form.seek(0)
        return form

    def get_overlay_canvas(self, canvas_page) -> io.BytesIO:
        data = None
        if canvas_page == 1:
            data = io.BytesIO()
            pdf = canvas.Canvas(data)
            pdf.setFont('Helvetica-Bold', 13)

            # Тут важны Y- положение
            if self.profile.photo_for_schengen:
                img = ImageReader(self.profile.photo_for_schengen.url)
                # if self.profile.photo_for_schengen.path[-4:] == '.pdf':
                #     try:
                #         # imgs = convert_from_path(self.profile.photo_for_schengen.path, 500)
                #
                #         jpg_filename = self.profile.photo_for_schengen.path.replace('.pdf', '.jpg')
                #         img.save(jpg_filename, 'JPEG')
                #         pdf.drawImage(jpg_filename, 340, 642, 112, 134)
                #     except:
                #         pass
                # else:
                try:
                    pdf.drawImage(img, 340, 642, 112, 134)
                except:
                    pass

            pdf.drawString(x=265, y=796, text=self.document_year)

            # Familienname
            pdf.drawString(x=180, y=728, text=self.profile.last_name)
            # Vorname
            pdf.drawString(x=180, y=700, text=self.profile.first_name)

            # Geschlecht
            #     weiblich
            if self.profile.gender == 'F':
                pdf.circle(180.8, 676, 4, fill=1)
            # pdf.ellipse(100, 100, 50, 50, fill=1)
            #       männlich
            if self.profile.gender == 'M':
                pdf.circle(232, 676, 4, fill=1)

            # Staatsangehörigkeit
            pdf.drawString(x=180, y=645, text=self.profile.nationality)
            # Geburtsdatum
            pdf.drawString(x=181, y=617, text=self.profile.bday.strftime('%d'))
            pdf.drawString(x=210, y=617, text=self.profile.bday.strftime('%m'))
            pdf.drawString(x=242, y=617, text=self.profile.bday.strftime('%Y'))

            # Geburtsort:
            pdf.drawString(x=355, y=617, text=self.profile.birth_place)
            # Straße mit Hausnummer
            pdf.drawString(x=301, y=576, text=self.profile.live_street_number_translit)
            # Postleitzahl und Ort
            postcode = '723500' if self.profile.live_city_en.lower() == 'osh' else '720000'
            custom_live_city = 'Osh' if self.profile.live_city_en.lower() == 'osh' else 'Bishkek'
            pdf.drawString(x=301, y=548, text=f'{postcode}, {custom_live_city}')
            # Land
            pdf.drawString(x=301, y=521, text='Kirgisistan')
            # Telefon mit Landesvorwahl und Ortsvorwahl
            pdf.drawString(x=301, y=493, text=f'+{self.profile.user.phone}') # городской телефон
            # Mobiltelefon mit Landesvorwahl
            pdf.drawString(x=301, y=466, text=f'+{self.profile.user.phone}')
            # Email - Adresse
            pdf.drawString(x=301, y=438, text=self.profile.user.email)
            # Skype – Adresse

            # TODO: add skype
            # pdf.drawString(x=301, y=410.5, text='ololo skype') # адрес skype
            ################### 1 ###########################
            # Sprachkenntnisse bitte selbst beurteilen und ankreuzen, wir testen später
            l = [180, 217, 253, 285.9, 323]
            d = {'1': 180.5, '2': 216.5, '3': 251.5, '4': 286.5, '5': 322.5}

            pdf.circle(d[self.profile.german], 283.5, 4.3, fill=1)
            # for i in range(5):
            #     pdf.circle(l[i], 282.9, 3.1, fill=1)

            ################### 2 ###########################
            pdf.circle(d[self.profile.english], 269.6, 4.3, fill=1)

            ################### 3 ###########################
            pdf.circle(d[self.profile.turkish], 256, 4.3, fill=1)

            ################### 4 ###########################
            pdf.circle(d[self.profile.russian], 242.3, 4.3, fill=1)

            pdf.drawString(x=70, y=253, text='Türkisch')
            pdf.drawString(x=70, y=239, text='Russisch')

            ################### ya nein #####################

            # Können Sie Fahhrad fahren? кататься на велосипеде

            answr = {'ja': {'x_cen': 286, 'y_cen': 186.5}, 'nein': {'x_cen': 368.5, 'y_cen': 187.5}}
            vybor = 'ja' if self.profile.bicycle_skill in ['ride_good', 'ride_bad'] else 'nein'
            pdf.circle(answr[vybor]['x_cen'], answr[vybor]['y_cen'], 7, fill=1)
            ################### ya nein #####################

            # Haben Sie Führerschein? водительские права

            answr2 = {'ja': 286.5, 'nein': 368.5, }
            vybor2 = 'ja' if self.profile.driver_license else 'nein'
            pdf.circle(answr2[vybor2], 159, 7, fill=1)
            ################### ya nein #####################

            # Haben Sie Fahrerfahrung? водительский стаж

            answr3 = {'ja': {'x_cen': 286.5, 'y_cen': 130}, 'nein': {'x_cen': 368.5, 'y_cen': 134.5}}
            vybor3 = 'ja' if self.profile.driving_experience else 'nein'
            pdf.circle(answr3[vybor3]['x_cen'], answr3[vybor3]['y_cen'], 7, fill=1)

            pdf.save()
            data.seek(0)

        elif canvas_page == 2:
            data = io.BytesIO()
            pdf = canvas.Canvas(data)
            pdf.setFont('Helvetica-Bold', 13)
            # Тут важны Y- положение
            # Daten zum Studium
            #   Student / in seit

            pdf.drawString(x=180, y=659, text=self.profile.study_start.strftime('%d'))
            pdf.drawString(x=210, y=659, text=self.profile.study_start.strftime('%m'))
            pdf.drawString(x=243, y=659, text=self.profile.study_start.strftime('%Y'))
            #   Student / in bis voraussichtlich

            pdf.drawString(x=442, y=659, text=self.profile.study_end.strftime('%d'))
            pdf.drawString(x=468, y=659, text=self.profile.study_end.strftime('%m'))
            pdf.drawString(x=497, y=659, text=self.profile.study_end.strftime('%Y'))
            # Studienfach
            pdf.drawString(x=180, y=631.3, text=self.profile.faculty.name_de)
            # Universität mit Adresse
            pdf.drawString(x=180, y=604, text=self.profile.university.name_de)
            pdf.drawString(x=180, y=576, text=self.profile.university.address)

            # Semesterferien
            # von
            pdf.drawString(x=216, y=548.3, text=self.profile.summer_holidays_start.strftime('%d'))
            pdf.drawString(x=246, y=548.3, text=self.profile.summer_holidays_start.strftime('%m'))
            pdf.drawString(x=276, y=548.3, text=self.profile.summer_holidays_start.strftime(self.document_year))
            # ----------------------------------------
            # bis
            pdf.drawString(x=357, y=548.3, text=self.profile.summer_holidays_end.strftime('%d'))
            pdf.drawString(x=387, y=548.3, text=self.profile.summer_holidays_end.strftime('%m'))
            pdf.drawString(x=417, y=548.3, text=self.profile.summer_holidays_end.strftime(self.document_year))

            # Gewünschter Beschäftigungszeitraum
            pdf.circle(95, 458.6, 7, fill=1)

            pdf.drawString(x=240, y=426, text=self.profile.summer_holidays_end.strftime(self.document_year))
            pdf.drawString(x=380, y=426, text=self.profile.summer_holidays_end.strftime(self.document_year))
            pdf.drawString(x=240, y=382, text=self.profile.summer_holidays_end.strftime(self.document_year))
            pdf.drawString(x=380, y=382, text=self.profile.summer_holidays_end.strftime(self.document_year))

            # TODO: нужны ли даты или достаточно галочки?
            # pdf.drawString(x=180, y=427.7, text='12')
            # pdf.drawString(x=210, y=427.7, text='12')
            #
            # pdf.drawString(x=321, y=427.7, text='12')
            # pdf.drawString(x=353, y=427.7, text='12')
            #
            # pdf.drawString(x=180, y=383.3, text='12')
            # pdf.drawString(x=210, y=383.3, text='12')
            #
            # pdf.drawString(x=321, y=383.3, text='12')
            # pdf.drawString(x=353, y=383.3, text='12')

            # Bisherige Beschäftigungen

            # first work place
            pdf.drawString(x=70, y=277, text=self.profile.start_date1.strftime('%d.%m.%Y'))
            pdf.drawString(x=140, y=277, text=self.profile.end_date1.strftime('%d.%m.%Y'))
            pdf.drawString(x=213, y=277, text=self.profile.position1)
            pdf.drawString(x=318, y=277, text=self.profile.company1)
            pdf.drawString(x=438, y=277, text=self.profile.country1)

            # second work place
            pdf.drawString(x=70, y=249, text=self.profile.start_date2.strftime('%d.%m.%Y'))
            pdf.drawString(x=140, y=249, text=self.profile.end_date2.strftime('%d.%m.%Y'))
            pdf.drawString(x=213, y=249, text=self.profile.position2)
            pdf.drawString(x=318, y=249, text=self.profile.company2)
            pdf.drawString(x=438, y=249, text=self.profile.country2)

            # third work place
            if self.profile.start_date3 and self.profile.end_date3 and\
                    self.profile.position3 and self.profile.company3 and self.profile.country3:
                pdf.drawString(x=70, y=221.5, text=self.profile.start_date3.strftime('%d.%m.%Y'))
                pdf.drawString(x=140, y=221.5, text=self.profile.end_date3.strftime('%d.%m.%Y'))
                pdf.drawString(x=213, y=221.5, text=self.profile.position3)
                pdf.drawString(x=318, y=221.5, text=self.profile.company3)
                pdf.drawString(x=438, y=221.5, text=self.profile.country3)


            # germany work experience
            # first work place
            germany_work_experience = {}
            if self.profile.country1 == 'Deutschland':
                germany_work_experience['first'] = {
                    'start_date': self.profile.start_date1.strftime('%d.%m.%Y'),
                    'end_date': self.profile.end_date1.strftime('%d.%m.%Y'),
                    'position': self.profile.position1,
                    'company': self.profile.company1,
                    'country': self.profile.country1,
                }

            if self.profile.country2 == 'Deutschland':
                germany_work_experience['second'] = {
                    'start_date': self.profile.start_date2.strftime('%d.%m.%Y'),
                    'end_date': self.profile.end_date2.strftime('%d.%m.%Y'),
                    'position': self.profile.position2,
                    'company': self.profile.company2,
                    'country': self.profile.country2,
                }

            if self.profile.country3 == 'Deutschland':
                germany_work_experience['third'] = {
                    'start_date': self.profile.start_date3.strftime('%d.%m.%Y'),
                    'end_date': self.profile.end_date3.strftime('%d.%m.%Y'),
                    'position': self.profile.position3,
                    'company': self.profile.company3,
                    'country': self.profile.country3,
                }

            i = 0
            if germany_work_experience:
                for k, v in germany_work_experience.items():
                    if i == 0:
                        pdf.drawString(x=70, y=112.5, text=v['start_date'])
                        pdf.drawString(x=140, y=112.5, text=v['end_date'])
                        pdf.drawString(x=213, y=112.5, text=v['position'])
                        pdf.drawString(x=318, y=112.5, text=v['company'])
                        pdf.drawString(x=438, y=112.5, text=v['country'])
                    elif i == 1:
                        pdf.drawString(x=70, y=84.2, text=v['start_date'])
                        pdf.drawString(x=140, y=84.2, text=v['end_date'])
                        pdf.drawString(x=213, y=84.2, text=v['position'])
                        pdf.drawString(x=318, y=84.2, text=v['company'])
                        pdf.drawString(x=438, y=84.2, text=v['country'])
                    elif i == 2:
                        pdf.drawString(x=70, y=57, text=v['start_date'])
                        pdf.drawString(x=140, y=57, text=v['end_date'])
                        pdf.drawString(x=213, y=57, text=v['position'])
                        pdf.drawString(x=318, y=57, text=v['company'])
                        pdf.drawString(x=438, y=57, text=v['country'])
                    i += 1
            else:
                pdf.drawString(x=70, y=112.5, text='Nein')


            # fourth work place
            # pdf.drawString(x=80, y=193.8, text='WARNER B CORP')

            # Waren Sie schon einmal in Deutschland?
            # TODO: блок вы были в Германии?
            # pdf.drawString(x=80, y=111, text='NETFLIX B CORP')
            # pdf.drawString(x=80, y=83.8, text='NETFLIX B CORP')
            # pdf.drawString(x=80, y=56, text='NETFLIX B CORP')

            pdf.save()
            data.seek(0)

        elif canvas_page == 3:
            data = io.BytesIO()
            pdf = canvas.Canvas(data)
            pdf.setFont('Helvetica-Bold', 13)
            # Тут важны Y- положение

            # Ich habe weitere Freunde, Verwandte und Bekannte, die auch über neue Jobs informiert werden wollen
            # TODO: нужен ли блок о друзья, которые хотят инфо о работе в Германии?
            # pdf.drawString(x=120, y=603, text='Steve')
            # pdf.drawString(x=230, y=603, text='American')
            # pdf.drawString(x=320, y=603, text='173')
            # pdf.drawString(x=360, y=603, text='captain@merica.com')
            #
            # pdf.drawString(x=120, y=576, text='Tony')
            # pdf.drawString(x=230, y=576, text='American')
            # pdf.drawString(x=320, y=576, text='45')
            # pdf.drawString(x=360, y=576, text='iron@man.com')
            #
            # pdf.drawString(x=120, y=548, text='Thor')
            # pdf.drawString(x=230, y=548, text='Asgardian')
            # pdf.drawString(x=320, y=548, text='1000')
            # pdf.drawString(x=360, y=548, text='thor@venger.com')
            #
            # pdf.drawString(x=120, y=521, text='Romanov')
            # pdf.drawString(x=230, y=521, text='Russian')
            # pdf.drawString(x=320, y=521, text='35')
            # pdf.drawString(x=360, y=521, text='black@widow.com')
            # #
            # pdf.drawString(x=120, y=491, text='Benner')
            # pdf.drawString(x=230, y=491, text='American')
            # pdf.drawString(x=320, y=491, text='40')
            # pdf.drawString(x=360, y=491, text='hulk@venger.com')

            # Mit meiner Unterschrift bestätige ich die Richtigkeit der Angaben:
            pdf.drawString(x=170, y=423, text=date.today().strftime('%d.%m.%Y'))

            pdf.save()
            data.seek(0)

        return data

    def save(self, form: io.BytesIO, filename: str):
        with open(filename, 'wb') as f:
            f.write(form.read())

    # fill first page
    def fill_first_page(self):
        canvas_page = 1
        template_path = self.TEMPLATE_PATH1
        filename = self.OUTPUT_PATH1
        self.run(canvas_page, template_path, filename)


    def fill_second_page(self):
        canvas_page = 2
        template_path = self.TEMPLATE_PATH2
        filename = self.OUTPUT_PATH2
        self.run(canvas_page, template_path, filename)

    def fill_third_page(self):
        canvas_page = 3
        template_path = self.TEMPLATE_PATH3
        filename = self.OUTPUT_PATH3
        self.run(canvas_page, template_path, filename)

    def merge_pages(self):
        paths = [self.OUTPUT_PATH1, self.OUTPUT_PATH2, self.OUTPUT_PATH3]
        pdf_writer = PdfFileWriter()

        for path in paths:
            pdf_reader = PdfFileReader(path)
            for page in range(pdf_reader.getNumPages()):
                # Add each page to the writer object
                pdf_writer.addPage(pdf_reader.getPage(page))

        # Write out the merged PDF
        with open(self.FINAL_PDF_PATH, 'wb') as out:
            pdf_writer.write(out)

        doclist = [self.OUTPUT_PATH1, self.OUTPUT_PATH2, self.OUTPUT_PATH3]
        for i in range(len(doclist)):
            os.remove(doclist[i])

        return self.FINAL_PDF_PATH

    def generate_jeople(self):
        if self.profile.is_confirmed and self.profile.is_admin_confirmed and self.profile.is_form_completed\
                and self.profile.summer_holidays_start and self.profile.summer_holidays_end:
            self.fill_first_page()
            self.fill_second_page()
            self.fill_third_page()
            return self.merge_pages()
        return None


class ContractGenerator:
    def __init__(self, profile):
        self.profile = profile

        self.dt = timezone.localtime()

        filename_pdf = f'{self.profile.last_name}-{self.profile.first_name}-{self.profile.user.email}-jeople-{self.dt}.pdf'

        self.BASE_PATH = os.path.dirname(os.path.dirname(os.path.abspath(__file__).replace('/applications', '')))
        self.COMMON_PATH = os.path.join(self.BASE_PATH, 'common_docs/')
        self.MEDIA_PATH = os.path.join(self.BASE_PATH, 'media/')

        self.DEFAUL_DIRECTORY = f'{self.COMMON_PATH}documents/'
        self.OUTPUT_DIRECTORY = f'{self.MEDIA_PATH}documents/'

    def get_month(self, num):
        months = {
            '01': 'января',
            '02': 'февраля',
            '03': 'марта',
            '04': 'апреля',
            '05': 'мая',
            '06': 'июня',
            '07': 'июля',
            '08': 'августа',
            '09': 'сентября',
            '10': 'октября',
            '11': 'ноября',
            '12': 'декабря',
        }
        return months[num]

    def generate_filename(self, document_type, admin=None, training_rate=None, employment_type=None):
        filename, output_filename = '', ''
        if document_type == 'training':
            name_by_rate = {'econom1': '-14', 'econom2': '-16', 'standard': '-18'}
            filename = f'{self.DEFAUL_DIRECTORY}training-with-completion-{admin}{name_by_rate[training_rate]}.docx'
            # filename = f'{self.DEFAUL_DIRECTORY}{admin}{name_by_rate[training_rate]}.docx'
            output_filename = f'{self.OUTPUT_DIRECTORY}{self.profile.last_name}-{self.profile.first_name}-{admin}{name_by_rate[training_rate]}.docx'
        elif document_type == 'employment':
            filename = f'{self.DEFAUL_DIRECTORY}employment-with-completion-{employment_type}.docx'
            # filename = f'{self.DEFAUL_DIRECTORY}employment-{employment_type}.docx'
            output_filename = f'{self.OUTPUT_DIRECTORY}{self.profile.last_name}-{self.profile.first_name}-employment-{employment_type}.docx'
        return filename, output_filename
    
    def generate_filename_agreement(self, document_type, document_name, admin=None, agreement_cost=None):
        
        if document_type == 'training':
            filename = f'{self.DEFAUL_DIRECTORY}{document_name}.docx'
            
            if agreement_cost:
                output_filename = f'{self.OUTPUT_DIRECTORY}{self.profile.last_name}-{self.profile.first_name}-training-{agreement_cost}.docx'
            else:
                output_filename = f'{self.OUTPUT_DIRECTORY}{self.profile.last_name}-{self.profile.first_name}-training-stable.docx'

        elif document_type == 'employment':
            filename = f'{self.DEFAUL_DIRECTORY}{document_name}.docx'
            output_filename = f'{self.OUTPUT_DIRECTORY}{self.profile.last_name}-{self.profile.first_name}-employment-{agreement_cost}.docx'

        elif document_type == 'act':
            filename = f'{self.DEFAUL_DIRECTORY}{document_name}.docx'
            output_filename = f'{self.OUTPUT_DIRECTORY}{self.profile.last_name}-{self.profile.first_name}-act.docx'

        return filename, output_filename

    def generate_output_filename(self):
        return f'{self.OUTPUT_DIRECTORY}{self.profile.last_name}-{self.profile.first_name}-training-.docx'


    # v.2 addition of documents
    def generate_training_agreement(self, admin, agreement_cost):
        filename, output_filename = self.generate_filename_agreement(document_type='training', document_name='training2020', agreement_cost=agreement_cost)
        document = Document(filename)
        site_settings = SiteSettings.load()

        cost_in_letters = 'Сумма прописью'

        style = document.styles['Normal']
        font = style.font
        font.name = 'Corbel'
        font.size = Pt(12)

        today = datetime.datetime.now()
        first_part = today.strftime("%d%m") + '20'

        second_part = site_settings.training_serial_number

        if second_part < 10:
            second_part = '00' + str(second_part)
        elif second_part < 100:
            second_part = '0' + str(second_part)

        document_number = first_part + second_part

        if not self.profile.agreement1_date and not self.profile.agreement1_number:            
            self.profile.agreement1_number = document_number
            self.profile.agreement1_date = datetime.date.today()
            self.profile.save(update_fields=['agreement1_number', 'agreement1_date'])

        number = document.paragraphs[0].text.replace('номер договора', document_number)
        document.paragraphs[0].text = number

        current_month_number = int(today.strftime("%m"))
        if current_month_number < 10:
            current_month = '0' + str(current_month_number)
        else:
            current_month = str(current_month_number)

        document_date = document.paragraphs[3].text.replace('дата первой выгрузки документа', today.strftime("%d") + " " + self.get_month(current_month) + " " + today.strftime("%Y"))
        document.paragraphs[3].text = document_date

        admin_name = document.paragraphs[5].text.replace('Талантбекова Гулина Талантбековна', admin.full_name)
        document.paragraphs[5].text = admin_name

        admin_is_male = admin.gender == ContractAdmin.MALE

        admin_suffix = document.paragraphs[5].text.replace('действующий/ая', 'действющий' if admin_is_male else 'действующая')
        document.paragraphs[5].text = admin_suffix

        patent_number = document.paragraphs[5].text.replace('№0854094 от 03.09.2019', f'{admin.patent_id}')
        document.paragraphs[5].text = patent_number

        patent_given_place = document.paragraphs[5].text.replace('УГНС по Ленинскому району', f'{admin.given_by}')
        document.paragraphs[5].text = patent_given_place

        admin_suffix1 = document.paragraphs[5].text.replace('расположенный/ая', 'расположенный' if admin_is_male else 'расположенная')
        document.paragraphs[5].text = admin_suffix1

        admin_suffix2 = document.paragraphs[5].text.replace('именуемый/ая', 'именуемый' if admin_is_male else 'именуемая')
        document.paragraphs[5].text = admin_suffix2

        student_is_male = self.profile.gender == 'M'

        student_suffix = document.paragraphs[5].text.replace('гражданин(ка)', 'гражданин' if student_is_male else 'гражданка')
        document.paragraphs[5].text = student_suffix

        student_full_name = document.paragraphs[5].text.replace('Фамилия Имя студента', f'{self.profile.full_name_ru}')
        document.paragraphs[5].text = student_full_name

        student_suffix1 = document.paragraphs[5].text.replace('прописанный/ая', 'прописанный' if student_is_male else 'прописанная')
        document.paragraphs[5].text = student_suffix1

        student_registered_address = document.paragraphs[5].text.replace('адрес по прописке студента', f'{self.profile.reg_address}')
        document.paragraphs[5].text = student_registered_address

        student_suffix2 = document.paragraphs[5].text.replace('проживающий/ая', 'проживающий' if student_is_male else 'проживающая')
        document.paragraphs[5].text = student_suffix2

        student_live_address = document.paragraphs[5].text.replace('фактический адрес студента', f'{self.profile.live_address}')
        document.paragraphs[5].text = student_live_address

        student_suffix3 = document.paragraphs[5].text.replace('Именуемый/ая', 'именуемый' if student_is_male else 'именуемая')
        document.paragraphs[5].text = student_suffix3


        document_cost = document.paragraphs[33].text.replace('14000 / 12000 /10000 (четырнадцать тысяч / двенадцать тысяч / десять тысяч)', 
                                                             str(agreement_cost.amount_in_digits) + ' ' + agreement_cost.amount_in_text)
        document.paragraphs[33].text = document_cost

        document_cost1 = document.paragraphs[42].text.replace('14000 / 12000 /10000 (четырнадцать тысяч / двенадцать тысяч / десять тысяч)', 
                                                              str(agreement_cost.amount_in_digits) + ' ' + agreement_cost.amount_in_text)
        document.paragraphs[42].text = document_cost1

        document_date = document.paragraphs[48].text.replace(
                                '«Дата первой выгрузки договора»',
                                today.strftime("%d") + " " + self.get_month(current_month) + " " + today.strftime("%Y")
                            )
        document.paragraphs[48].text = document_date

        admin_full_name = document.tables[0].rows[1].cells[0].text.replace('Талантбекова Гулина', f'{admin.full_name}')
        document.tables[0].rows[1].cells[0].text = admin_full_name

        student_full_name = document.tables[0].rows[1].cells[1].text.replace('Фамилия Имя студента', f'{self.profile.full_name_ru}')
        document.tables[0].rows[1].cells[1].text = student_full_name

        for paragraph in document.paragraphs:
            paragraph.style = document.styles['Normal']
        for run in paragraph.runs:
            run.font.size = Pt(11)

        document.save(output_filename)

        return output_filename

    def generate_act_for_training(self, admin):

        filename, output_filename = self.generate_filename_agreement(document_name='act2020', document_type='act')

        document = Document(filename)
        site_settings = SiteSettings.load()
        agreement_date = self.profile.agreement1_date

        style = document.styles['Normal']
        font = style.font
        font.name = 'Corbel'
        font.size = Pt(12)

        month_number = int(self.profile.agreement1_date.strftime("%m"))
        if month_number < 10:
            month_string = '0' + str(month_number)
        else:
            month_string = str(month_number)

        document_number = self.profile.agreement1_number

        number_and_date_string = document_number + ' от «' + agreement_date.strftime("%d") + "» " + self.get_month(month_string) + " " + agreement_date.strftime("%Y")
        agreement_number_and_date = document.paragraphs[1].text.replace('_______ (номер Договора) от «___»______________ 20__г. Дата первой выгрузки договора', number_and_date_string)
        document.paragraphs[1].text = agreement_number_and_date

        admin_name = document.paragraphs[5].text.replace('Талантбекова Гулина Талантбековна патентщик', 'ИП ' + admin.full_name)
        document.paragraphs[5].text = admin_name

        admin_is_male = admin.gender == ContractAdmin.MALE

        admin_suffix = document.paragraphs[5].text.replace('действующий/ая', 'действующий' if admin_is_male else 'действующая')
        document.paragraphs[5].text = admin_suffix

        patent_number = document.paragraphs[5].text.replace('№0854094 от 03.09.2019', f'{admin.patent_id}')
        document.paragraphs[5].text = patent_number

        patent_given_place = document.paragraphs[5].text.replace('УГНС по Ленинскому району', f'{admin.given_by}')
        document.paragraphs[5].text = patent_given_place

        admin_suffix1 = document.paragraphs[5].text.replace('расположенный/ая', 'расположенный' if admin_is_male else 'расположенная')
        document.paragraphs[5].text = admin_suffix1

        admin_suffix2 = document.paragraphs[5].text.replace('именуемый/ая', 'именуемый' if admin_is_male else 'именуемая')
        document.paragraphs[5].text = admin_suffix2

        student_is_male = self.profile.gender == 'M'

        student_suffix = document.paragraphs[5].text.replace('гражданин/ка', 'гражданин' if student_is_male else 'гражданка')
        document.paragraphs[5].text = student_suffix

        student_full_name = document.paragraphs[5].text.replace('Фамилия имя студента', self.profile.full_name_ru)
        document.paragraphs[5].text = student_full_name

        student_suffix1 = document.paragraphs[5].text.replace('прописанный/ая', 'прописанный' if student_is_male else 'прописанная')
        document.paragraphs[5].text = student_suffix1

        student_registered_address = document.paragraphs[5].text.replace('адрес по прописке', self.profile.reg_address)
        document.paragraphs[5].text = student_registered_address

        student_suffix2 = document.paragraphs[5].text.replace('проживающий/ая', 'проживающий' if student_is_male else 'проживающая')
        document.paragraphs[5].text = student_suffix2

        student_live_address = document.paragraphs[5].text.replace('фактический адрес', self.profile.live_address)
        document.paragraphs[5].text = student_live_address

        student_suffix3 = document.paragraphs[5].text.replace('именуемый/ая', 'именуемый' if student_is_male else 'именуемая')
        document.paragraphs[5].text = student_suffix3

        number_and_date_string = document_number + " " + agreement_date.strftime("%d") + " " + self.get_month(month_string) + " " + agreement_date.strftime("%Y")
        agreement_number_and_date = document.paragraphs[5].text.replace('номер и дата первой выгрузки Договора', number_and_date_string)
        document.paragraphs[5].text = agreement_number_and_date

        admin_full_name = document.tables[1].rows[1].cells[0].text.replace('Талантбекова Гулина', admin.full_name)
        document.tables[1].rows[1].cells[0].text = admin_full_name

        student_full_name = document.tables[1].rows[1].cells[1].text.replace('   Фамилия имя студента', self.profile.full_name_ru)
        document.tables[1].rows[1].cells[1].text = student_full_name

        for paragraph in document.paragraphs:
            paragraph.style = document.styles['Normal']
        for run in paragraph.runs:
            run.font.size = Pt(11)

        document.save(output_filename)

        return output_filename
    
    def generate_act_for_employment(self):

        filename, output_filename = self.generate_filename_agreement(document_name='employment-act-2020', document_type='act')

        document = Document(filename)
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        site_settings = SiteSettings.load()

        agreement_date = self.profile.agreement2_date

        document_number = self.profile.agreement2_number

        month_number = int(agreement_date.strftime("%m"))
        if month_number < 10:
            month_string = '0' + str(month_number)
        else:
            month_string = str(month_number)

        number_and_date_string = document_number + ' от «' + agreement_date.strftime("%d") + "» " + self.get_month(month_string) + " " + agreement_date.strftime("%Y")
        agreement_number_and_date = document.paragraphs[2].text.replace('_______ от «___»________ 20__г.  Номер первого договора и Дата первой выгрузки договора', number_and_date_string)
        document.paragraphs[2].text = agreement_number_and_date

        student_is_male = self.profile.gender == 'M'
        student_suffix = document.paragraphs[9].text.replace('гражданин/ка', 'гражданин' if student_is_male else 'гражданка')
        document.paragraphs[9].text = student_suffix

        student_full_name = document.paragraphs[9].text.replace('Фамилия имя студента', self.profile.full_name_ru)
        document.paragraphs[9].text = student_full_name

        student_suffix1 = document.paragraphs[9].text.replace('прописанный/ая', 'прописаннный' if student_is_male else 'прописанная')
        document.paragraphs[9].text = student_suffix1

        student_reg_address = document.paragraphs[9].text.replace('адрес прописки студента', self.profile.reg_address)
        document.paragraphs[9].text = student_reg_address

        student_suffix2 = document.paragraphs[9].text.replace('проживающий/ая', 'проживающий' if student_is_male else 'проживающая')
        document.paragraphs[9].text = student_suffix2

        student_address = document.paragraphs[9].text.replace('Фактический адрес студента', self.profile.live_address)
        document.paragraphs[9].text = student_address

        student_suffix3 = document.paragraphs[9].text.replace('действующий/ая', 'действующий' if student_is_male else 'действующая')
        document.paragraphs[9].text = student_suffix3

        student_passport_number = document.paragraphs[9].text.replace('номер паспорта ID или AN', self.profile.passport_number)
        document.paragraphs[9].text = student_passport_number

        student_suffix4 = document.paragraphs[9].text.replace('именуемый/ая', 'именуемый' if student_is_male else 'именуемая')
        document.paragraphs[9].text = student_suffix4

        agreement_number = document.paragraphs[9].text.replace('_______ от «___»________ 20__г.  Номер первого договора и Дата первой выгрузки договора,', number_and_date_string)
        document.paragraphs[9].text = agreement_number

        student_full_name = document.tables[0].rows[1].cells[1].text.replace('Фамилия Имя Студента ', self.profile.full_name_ru)
        document.tables[0].rows[1].cells[1].text = student_full_name

        document.save(output_filename)

        return output_filename

    def generate_act_for_stable(self, admin):
        
        filename, output_filename = self.generate_filename_agreement(document_name='training-act-6000', document_type='act')

        document = Document(filename)
        site_settings = SiteSettings.load()
        today = datetime.datetime.now()

        document_number = self.profile.agreement3_number
        agreement_date = self.profile.agreement3_date

        month_number = int(agreement_date.strftime("%m"))
        if month_number < 10:
            month_string = '0' + str(month_number)
        else:
            month_string = str(month_number)

        number_and_date_string = document_number + ' от «' + agreement_date.strftime("%d") + "» " + self.get_month(month_string) + " " + agreement_date.strftime("%Y")
        agreement_number_and_date = document.paragraphs[2].text.replace('_______ (номер Договора )  от «___»______________ 20__г. Дата первой выгрузки договора', number_and_date_string)
        document.paragraphs[2].text = agreement_number_and_date

        patent_owner_name = document.paragraphs[9].text.replace('Талантбекова Гулина 	Талантбековна патентщик', admin.full_name)
        document.paragraphs[9].text = patent_owner_name

        admin_is_male = admin.gender == 'M'
        admin_suffix = document.paragraphs[9].text.replace('действующий/ая', 'действующий' if admin_is_male else 'действующая')
        document.paragraphs[9].text = admin_suffix

        patent_number = document.paragraphs[9].text.replace('№0854094 от 03.09.2019', f'{admin.patent_id}')
        document.paragraphs[9].text = patent_number

        patent_given_place = document.paragraphs[9].text.replace('УГНС по Ленинскому району', f'{admin.given_by}')
        document.paragraphs[9].text = patent_given_place

        admin_suffix1 = document.paragraphs[9].text.replace('расположенный/ая', 'расположенный' if admin_is_male else 'расположенная')
        document.paragraphs[9].text = admin_suffix1

        admin_suffix2 = document.paragraphs[9].text.replace('именуемый/ая', 'именуемый' if admin_is_male else 'именуемая')
        document.paragraphs[9].text = admin_suffix2

        student_is_male = self.profile.gender == 'M'

        student_suffix = document.paragraphs[9].text.replace('гражданин/ка', 'гражданин' if student_is_male else 'гражданка')
        document.paragraphs[9].text = student_suffix

        student_full_name = document.paragraphs[9].text.replace('Фамилия имя студента', self.profile.full_name_ru)
        document.paragraphs[9].text = student_full_name

        student_suffix1 = document.paragraphs[9].text.replace('прописанный/ая', 'прописанный' if student_is_male else 'прописанная')
        document.paragraphs[9].text = student_suffix1

        student_registered_address = document.paragraphs[9].text.replace('адрес по прописке', self.profile.reg_address)
        document.paragraphs[9].text = student_registered_address

        student_suffix2 = document.paragraphs[9].text.replace('проживающий/ая', 'проживающий' if student_is_male else 'проживающая')
        document.paragraphs[9].text = student_suffix2

        student_live_address = document.paragraphs[9].text.replace('фактический адрес', self.profile.live_address)
        document.paragraphs[9].text = student_live_address

        student_suffix3 = document.paragraphs[9].text.replace('именуемый/ая', 'именуемый' if student_is_male else 'именуемая')
        document.paragraphs[9].text = student_suffix3

        agreement_number_and_date = document.paragraphs[9].text.replace('номер и дата первой выгрузки Договора', number_and_date_string)
        document.paragraphs[9].text = agreement_number_and_date

        admin_full_name = document.tables[1].rows[1].cells[0].text.replace('Талантбекова Гулина', admin.full_name)
        document.tables[1].rows[1].cells[0].text = admin_full_name

        student_full_name = document.tables[1].rows[1].cells[1].text.replace('        Фамилия имя студента', self.profile.full_name_ru)
        document.tables[1].rows[1].cells[1].text = student_full_name

        for paragraph in document.paragraphs:
            paragraph.style = document.styles['Normal']
        for run in paragraph.runs:
            run.font.size = Pt(11)

        document.save(output_filename)

        return output_filename

    def replace_text_in_paragraph(self, paragraph_number: int, existing_text: str, replacement_text: str) -> None:

        replacement = self.document.paragraphs[paragraph_number].text.replace(existing_text, replacement_text)
        self.document.paragraphs[paragraph_number].text = replacement

    def change_style_of_document(self, font_name: str, font_size: int) -> None:
        style = self.document.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)

    def generate_employment_agreement(self, agreement_cost):

        filename, output_filename = self.generate_filename_agreement(document_type='employment', document_name='employment2020', agreement_cost=agreement_cost)

        self.document = Document(filename)
        site_settings = SiteSettings.load()
        self.change_style_of_document('Arial', 12)

        today = datetime.datetime.now()
        first_part = today.strftime("%d%m") + '20'

        second_part = site_settings.employment_serial_number

        if second_part < 10:
            second_part = '00' + str(second_part)
        elif second_part < 100:
            second_part = '0' + str(second_part)

        document_number = first_part + second_part

        if not self.profile.agreement2_date and not self.profile.agreement2_number:
            self.profile.agreement2_date = datetime.date.today()
            self.profile.agreement2_number = document_number

            self.profile.save(update_fields=['agreement2_date', 'agreement2_number'])

        number = self.document.paragraphs[0].text.replace('номер договора', document_number)
        self.document.paragraphs[0].text = number

        current_month_number = int(self.profile.agreement2_date .strftime("%m"))
        if current_month_number < 10:
            current_month = '0' + str(current_month_number)
        else:
            current_month = str(current_month_number)

        self.replace_text_in_paragraph(2, 'Дата первой выгрузки договора', self.profile.agreement2_date.strftime("%d") + " " + self.get_month(current_month) + " " + self.profile.agreement2_date.strftime("%Y"))

        student_is_male = self.profile.gender == 'M'

        self.replace_text_in_paragraph(4, 'гражданин/ка', 'гражданин' if student_is_male else "гражданка")

        self.replace_text_in_paragraph(4, 'имеющий/ая', 'имеющий' if student_is_male else 'имеющая')
        
        self.replace_text_in_paragraph(4, 'гражданина/ки', 'гражданинa' if student_is_male else 'гражданки')

        self.replace_text_in_paragraph(4, 'обладающий/ая', 'обладающий' if student_is_male else 'обладающая')

        self.replace_text_in_paragraph(4, 'являющийся/аяся', 'являющийся' if student_is_male else "являющаяся")

        self.replace_text_in_paragraph(4, 'Фамилия имя студента   ', f'{self.profile.full_name_ru}')

        self.replace_text_in_paragraph(4, 'прописанный/ая', 'прописанный' if student_is_male else 'прописанная')

        self.replace_text_in_paragraph(4, 'адрес прописки студента', f'{self.profile.reg_address}')

        self.replace_text_in_paragraph(4, 'проживающий/ая', 'проживающий' if student_is_male else 'проживающая')

        self.replace_text_in_paragraph(4, 'Фактический адрес студента', f'{self.profile.live_address}')
        
        self.replace_text_in_paragraph(4, 'действующий/ая', "действующий" if student_is_male else "действующая")

        self.replace_text_in_paragraph(4, 'номер паспорта ID или AN', f'{self.profile.passport_number}')

        self.replace_text_in_paragraph(4, 'именуемый/ая', 'именуемый' if student_is_male else 'именуемая')

        number_in_letters = {
            '1000': "одна тысяча сомов",
            "2000": "две тысячи сомов"
        }

        self.replace_text_in_paragraph(68, '2000 / 1000 сомов (две тысячи / одна тысяча сомов)', agreement_cost + " " + number_in_letters[agreement_cost])

        current_month_number = int(today.strftime("%m"))
        if current_month_number < 10:
            current_month = '0' + str(current_month_number)
        else:
            current_month = str(current_month_number)

        self.replace_text_in_paragraph(77, '"день" месяц год (дата первой выгрузки договора)', self.profile.agreement2_date.strftime("%d") + " " + self.get_month(current_month) + " " + self.profile.agreement2_date.strftime("%Y"))

        new_text = self.document.tables[0].rows[1].cells[1].text.replace('Фамилия имя студента', self.profile.full_name_ru)
        self.document.tables[0].rows[1].cells[1].text = new_text

        self.document.save(output_filename)

        return output_filename


    def generate_training_unchange(self, admin):
        filename, output_filename = self.generate_filename_agreement(document_type='training', document_name='training6000')

        self.document = Document(filename)
        site_settings = SiteSettings.load()
        self.change_style_of_document('Times New Roman', 11)

        today = datetime.datetime.now()
        first_part = today.strftime("%d%m") + '20'

        second_part = site_settings.training_serial_number

        if second_part < 10:
            second_part = '00' + str(second_part)
        elif second_part < 100:
            second_part = '0' + str(second_part)

        document_number = first_part + second_part

        if not self.profile.agreement3_date and not self.profile.agreement3_number:

            self.profile.agreement3_number = document_number
            self.profile.agreement3_date = datetime.date.today()
            self.profile.save(update_fields=['agreement3_number', 'agreement3_date'])

        self.replace_text_in_paragraph(0, 'номер договора', document_number)

        current_month_number = int(self.profile.agreement3_date.strftime("%m"))
        if current_month_number < 10:
            current_month = '0' + str(current_month_number)
        else:
            current_month = str(current_month_number)

        self.replace_text_in_paragraph(3, 'дата первой выгрузки документа', self.profile.agreement3_date.strftime("%d") + " " + self.get_month(current_month) + " " + self.profile.agreement3_date.strftime("%Y"))

        self.replace_text_in_paragraph(5, 'Талантбекова Гулина Талантбековна', admin.full_name)

        admin_is_male = admin.gender == 'M'

        self.replace_text_in_paragraph(5, 'действующий/ая', 'действующий' if admin_is_male else 'действующая')

        self.replace_text_in_paragraph(5, '№0854094 от 03.09.2019', f'{admin.patent_id}' + ' от ' + f'{admin.patent_date}')

        self.replace_text_in_paragraph(5, 'УГНС по Ленинскому району', f'{admin.given_by}')

        self.replace_text_in_paragraph(5, 'расположенный/ая', 'расположенный' if admin_is_male else 'расположенная')

        self.replace_text_in_paragraph(5, 'именуемый/ая', 'именуемый' if admin_is_male else "именуемая")

        student_is_male = self.profile.gender == 'M'

        self.replace_text_in_paragraph(5, 'гражданин(ка)', 'гражданин' if student_is_male else 'гражданка')

        self.replace_text_in_paragraph(5, 'Фамилия Имя студента', self.profile.full_name_ru)

        self.replace_text_in_paragraph(5, 'прописанный/ая', 'прописанный' if student_is_male else 'прописанная')

        self.replace_text_in_paragraph(5, 'адрес по прописке студента', self.profile.reg_address)

        self.replace_text_in_paragraph(5, 'проживающий/ая', 'проживающий' if student_is_male else 'проживающая')

        self.replace_text_in_paragraph(5, 'фактический адрес студента', self.profile.live_address)

        self.replace_text_in_paragraph(5, 'Именуемый/ая', "именуемый" if student_is_male else 'именуемая')

        self.replace_text_in_paragraph(47, '«Дата первой выгрузки договора»', today.strftime("%d") + " " + self.get_month(current_month) + " " + today.strftime("%Y"))

        admin_name = self.document.tables[0].rows[1].cells[0].text.replace('Талантбекова Гулина ', admin.full_name)
        self.document.tables[0].rows[1].cells[0].text = admin_name

        student_name = self.document.tables[0].rows[1].cells[1].text.replace('Фамилия Имя студента', self.profile.full_name_ru)
        self.document.tables[0].rows[1].cells[1].text = student_name

        self.replace_text_in_paragraph(118, '№ _______ (номер Договора )  от «___»______________ 20__г. Дата первой выгрузки договора',
                                       '№ ' + document_number + ' от ' + today.strftime("%d") + " " + self.get_month(current_month) + " " + today.strftime("%Y") + "г.")

        self.replace_text_in_paragraph(125, 'Талантбекова Гулина 	Талантбековна патентщик', admin.full_name)

        self.replace_text_in_paragraph(125, 'действующий/ая', 'действующий' if admin_is_male else 'действующая')

        self.replace_text_in_paragraph(125, '№0854094 от 03.09.2019', f'{admin.patent_id}' + ' от ' + f'{admin.patent_date}')

        self.replace_text_in_paragraph(125, 'УГНС по Ленинскому району', f'{admin.given_by}')

        self.replace_text_in_paragraph(125, 'расположенный/ая', 'расположенный' if admin_is_male else 'расположенная')

        self.replace_text_in_paragraph(125, 'именуемый/ая', 'именуемый' if admin_is_male else "именуемая")

        student_is_male = self.profile.gender == 'M'

        self.replace_text_in_paragraph(125, 'гражданин/ка', 'гражданин' if student_is_male else 'гражданка')

        self.replace_text_in_paragraph(125, 'Фамилия имя студента', self.profile.full_name_ru)

        self.replace_text_in_paragraph(125, 'прописанный/ая', 'прописанный' if student_is_male else 'прописанная')

        self.replace_text_in_paragraph(125, 'адрес по прописке', self.profile.reg_address)

        self.replace_text_in_paragraph(125, 'проживающий/ая', 'проживающий' if student_is_male else 'проживающая')

        self.replace_text_in_paragraph(125, 'фактический адрес', self.profile.live_address)

        self.replace_text_in_paragraph(125, 'Именуемый/ая', "именуемый" if student_is_male else 'именуемая')

        self.replace_text_in_paragraph(125, 'номер и дата первой выгрузки Договора', 
                                       document_number + ' от ' + self.profile.agreement3_date.strftime("%d") + " " + self.get_month(current_month) + " " + self.profile.agreement3_date.strftime("%Y") + "г.")

        admin_name = self.document.tables[2].rows[1].cells[0].text.replace('Талантбекова Гулина', admin.full_name)
        self.document.tables[2].rows[1].cells[0].text = admin_name

        student_name = self.document.tables[2].rows[1].cells[1].text.replace('Фамилия имя студента', self.profile.full_name_ru)
        self.document.tables[2].rows[1].cells[1].text = student_name

        for paragraph in self.document.paragraphs:
            paragraph.style = self.document.styles['Normal']
        for run in paragraph.runs:
            run.font.size = Pt(12)

        self.document.save(output_filename)

        return output_filename
    
    def generate_close_training(self, admin):
        filename = f'{self.DEFAUL_DIRECTORY}close2020.docx'
        output_filename = f'{self.OUTPUT_DIRECTORY}close--{self.profile.full_name}.docx'

        document = Document(filename)

        style = document.styles['Normal']
        font = style.font
        font.name = 'Corbel'
        font.size = Pt(12)

        newtitle = document.paragraphs[0].text.replace('дата  договора тренинги', self.profile.agreement1_date.strftime('%d/%m/%Y'))
        document.paragraphs[0].text = newtitle

        today = datetime.datetime.now()

        current_month_number = int(today.strftime("%m"))
        if current_month_number < 10:
            current_month = '0' + str(current_month_number)
        else:
            current_month = str(current_month_number)

        date = document.paragraphs[3].text.replace('дата выгрузки данного соглашения ', today.strftime("%d") + " " + self.get_month(current_month) + " " + today.strftime("%Y"))
        document.paragraphs[3].text = date

        pat = document.paragraphs[5].text.replace('Талантбекова Гулина Талантбековна (патентщики)', admin.full_name)
        document.paragraphs[5].text = pat

        admin_is_male = admin.gender == 'M'

        admin_suffix = document.paragraphs[5].text.replace('именуемый/ая', 'именуемый' if admin_is_male else 'именуемая')
        document.paragraphs[5].text = admin_suffix

        admin_name = document.paragraphs[5].text.replace('Талантбекова Гулина Талантбековна', admin.full_name)
        document.paragraphs[5].text = admin_name

        admin_suffix2 = document.paragraphs[5].text.replace('действующий/ая', 'действующий' if admin_is_male else 'действующая')
        document.paragraphs[5].text = admin_suffix2

        admin_patent = document.paragraphs[5].text.replace('№0854094 от 03.09.2019', admin.patent_id)
        document.paragraphs[5].text = admin_patent

        patent_address = document.paragraphs[5].text.replace('УГНС по Ленинскому району', admin.given_by)
        document.paragraphs[5].text = patent_address

        student_name = document.paragraphs[5].text.replace('фамилия и имя студента', self.profile.full_name_ru)
        document.paragraphs[5].text = student_name

        student_is_male = self.profile.gender == 'M'

        student_suffix = document.paragraphs[5].text.replace('Именуемый/ая', 'именуемый' if student_is_male else 'именуемая')
        document.paragraphs[5].text = student_suffix

        pass_number = document.paragraphs[5].text.replace('номер паспорта ID или AN', self.profile.passport_number)
        document.paragraphs[5].text = pass_number

        agr_num = document.paragraphs[7].text.replace('номер договора', self.profile.agreement1_number)
        document.paragraphs[7].text = agr_num

        agr_date = document.paragraphs[7].text.replace('дата договора', self.profile.agreement1_date.strftime('%d/%m/%Y'))
        document.paragraphs[7].text = agr_date

        date = document.paragraphs[10].text.replace('дата выгрузки данного соглашения', today.strftime("%d") + " " + self.get_month(current_month) + " " + today.strftime("%Y"))
        document.paragraphs[10].text = date

        agr_num = document.paragraphs[11].text.replace('номер договора', self.profile.agreement1_number)
        document.paragraphs[11].text = agr_num

        agr_date = document.paragraphs[11].text.replace('дата договора', self.profile.agreement1_date.strftime('%d/%m/%Y'))
        document.paragraphs[11].text = agr_date

        fam_name = document.paragraphs[16].text.replace('фамилия и имя патентщика', admin.full_name)
        document.paragraphs[16].text = fam_name

        student_name = document.paragraphs[19].text.replace('фамилия и имя студента', self.profile.full_name_ru)
        document.paragraphs[19].text = student_name

        fam_name = document.paragraphs[25].text.replace('фамилия и имя патентщика', admin.full_name)
        document.paragraphs[25].text = fam_name

        student_name = document.paragraphs[30].text.replace('Фамилия имя студента', self.profile.full_name_ru)
        document.paragraphs[30].text = student_name

        for paragraph in document.paragraphs:
            paragraph.style = document.styles['Normal']
        for run in paragraph.runs:
            run.font.size = Pt(11)

        document.save(output_filename)

        return output_filename

    def generate_close_employment(self):
        filename = f'{self.DEFAUL_DIRECTORY}close_employment.docx'
        output_filename = f'{self.OUTPUT_DIRECTORY}close_employment--{self.profile.full_name}.docx'

        document = Document(filename)

        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        newtitle = document.paragraphs[0].text.replace('дата первого договора', self.profile.agreement1_date.strftime('%d/%m/%Y'))
        document.paragraphs[0].text = newtitle

        date = document.paragraphs[3].text.replace('дата  последней выгрузки данного расторжения в формате « ДД » месяц год',
                                                   self.profile.agreement2_date.strftime('%d/%m/%Y'))
        document.paragraphs[3].text = date

        name_st = document.paragraphs[5].text.replace('Фамилия и Имя студента', self.profile.full_name_ru)
        document.paragraphs[5].text = name_st

        student_is_male = self.profile.gender == 'M'

        student_suffix = document.paragraphs[5].text.replace('именуемый/ая', 'именуемый' if student_is_male else 'именуемая')
        document.paragraphs[5].text = student_suffix

        student_suffix2 = document.paragraphs[5].text.replace('действующий/ая', 'действующий' if student_is_male else 'действующая')
        document.paragraphs[5].text = student_suffix2

        pass_num = document.paragraphs[5].text.replace('номер паспорта ID или AN', self.profile.passport_number)
        document.paragraphs[5].text = pass_num

        month_number = int(self.profile.agreement2_date.strftime("%m"))
        if month_number < 10:
            month = '0' + str(month_number)
        else:
            month = str(month_number)

        date = document.paragraphs[7].text.replace('номер договора от дата первого договора',
        self.profile.agreement2_number + "от" + self.profile.agreement2_date.strftime("%d") + " " + self.get_month(month) + " " + self.profile.agreement2_date.strftime("%Y"))
        document.paragraphs[7].text = date

        today = datetime.datetime.now()

        month_number = int(today.strftime("%m"))
        if month_number < 10:
            month2 = '0' + str(month_number)
        else:
            month2 = str(month_number)

        date2 = document.paragraphs[10].text.replace('дата выгрузки данного расторжения', today.strftime("%d") + " " + self.get_month(month2) + " " + today.strftime("%Y"))
        document.paragraphs[10].text = date2

        date3 = document.paragraphs[11].text.replace('дата первого договора', self.profile.agreement2_date.strftime("%d") + " " + self.get_month(month) + " " + self.profile.agreement2_date.strftime("%Y"))
        document.paragraphs[11].text = date3

        name_st = document.paragraphs[19].text.replace('Фамилия и имя студента', self.profile.full_name_ru)
        document.paragraphs[19].text = name_st

        address = document.paragraphs[20].text.replace('Фактический адрес студента', self.profile.live_address)
        document.paragraphs[20].text = address

        name_st2 = document.paragraphs[34].text.replace('Фамилия имя студента', self.profile.full_name_ru)
        document.paragraphs[34].text = name_st2

        for paragraph in document.paragraphs:
            paragraph.style = document.styles['Normal']
        for run in paragraph.runs:
            run.font.size = Pt(11)
        
        document.save(output_filename)

        return output_filename

    def generate_close_training_stable(self, admin):
        filename = f'{self.DEFAUL_DIRECTORY}close6000.docx'
        output_filename = f'{self.OUTPUT_DIRECTORY}close_6000--{self.profile.full_name}.docx'

        document = Document(filename)
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        month_number = int(self.profile.agreement3_date.strftime("%m"))
        if month_number < 10:
            month = '0' + str(month_number)
        else:
            month = str(month_number)

        newtitle = document.paragraphs[0].text.replace('дата  договора тренинги', self.profile.agreement3_date.strftime("%d") + " " + self.get_month(month) + " " + self.profile.agreement3_date.strftime("%Y"))
        document.paragraphs[0].text = newtitle

        today = datetime.datetime.now()

        month_number = int(today.strftime("%m"))
        if month_number < 10:
            month2 = '0' + str(month_number)
        else:
            month2 = str(month_number)

        date1 = document.paragraphs[3].text.replace('дата выгрузки данного соглашения', today.strftime("%d") + " " + self.get_month(month2) + " " + today.strftime("%Y"))
        document.paragraphs[3].text = date1

        admin_name = document.paragraphs[5].text.replace('Талантбекова Гулина Талантбековна (патентщики)', admin.full_name)
        document.paragraphs[5].text = admin_name

        admin_is_male = admin.gender == 'M'
        admin_suffix = document.paragraphs[5].text.replace('именуемый/ая', 'именуемый' if admin_is_male else 'именуемая')
        document.paragraphs[5].text = admin_suffix

        admin_name = document.paragraphs[5].text.replace('Талантбекова Гулина Талантбековна', admin.full_name)
        document.paragraphs[5].text = admin_name

        admin_suffix2 = document.paragraphs[5].text.replace('действующий/ая', 'действующий' if admin_is_male else 'действующая')
        document.paragraphs[5].text = admin_suffix2

        pat = document.paragraphs[5].text.replace('№0854094 от 03.09.2019', admin.patent_id)
        document.paragraphs[5].text = pat

        address = document.paragraphs[5].text.replace('УГНС по Ленинскому району', admin.given_by)
        document.paragraphs[5].text = address

        name = document.paragraphs[5].text.replace('фамилия и имя студента', self.profile.full_name_ru)
        document.paragraphs[5].text = name

        student_is_male = self.profile.gender == 'M'
        student_suffix = document.paragraphs[5].text.replace('Именуемый/ая', 'именуемый' if student_is_male else 'именуемая')
        document.paragraphs[5].text = student_suffix

        student_suffix2 = document.paragraphs[5].text.replace('Действующий/ая', 'действующий' if student_is_male else 'действующая')
        document.paragraphs[5].text = student_suffix2

        student_pass = document.paragraphs[5].text.replace('номер паспорта ID или AN', self.profile.passport_number)
        document.paragraphs[5].text = student_pass

        agr_number = document.paragraphs[7].text.replace('N номер договора', '№ ' + self.profile.agreement3_number)
        document.paragraphs[7].text = agr_number

        date3 = document.paragraphs[7].text.replace('дата договора', self.profile.agreement3_date.strftime("%d") + " " + self.get_month(month) + " " + self.profile.agreement3_date.strftime("%Y"))
        document.paragraphs[7].text = date3

        date4 = document.paragraphs[10].text.replace('дата выгрузки данного соглашения', today.strftime("%d") + " " + self.get_month(month2) + " " + today.strftime("%Y"))
        document.paragraphs[10].text = date4

        number1 = document.paragraphs[11].text.replace('номер договора', self.profile.agreement3_number)
        document.paragraphs[11].text = number1

        date3 = document.paragraphs[11].text.replace('дата договора', self.profile.agreement3_date.strftime("%d") + " " + self.get_month(month) + " " + self.profile.agreement3_date.strftime("%Y"))
        document.paragraphs[11].text = date3
        
        admin_name = document.paragraphs[16].text.replace('фамилия и имя патентщика', admin.full_name)
        document.paragraphs[16].text = admin_name

        name = document.paragraphs[19].text.replace('фамилия и имя студента', self.profile.full_name_ru)
        document.paragraphs[19].text = name

        address = document.paragraphs[20].text.replace('Фактисечкий адрес студента', self.profile.live_address)
        document.paragraphs[20].text = address

        admin_name = document.paragraphs[25].text.replace('фамилия и имя патентщика', admin.full_name)
        document.paragraphs[25].text = admin_name

        name = document.paragraphs[30].text.replace('Фамилия имя студента', self.profile.full_name_ru)
        document.paragraphs[30].text = name

        for paragraph in document.paragraphs:
            paragraph.style = document.styles['Normal']
        for run in paragraph.runs:
            run.font.size = Pt(11)
        
        document.save(output_filename)

        return output_filename


class CVGenerator:
    def __init__(self, profile):
        self.profile = profile

        self.dt = timezone.localtime()
        self.today = date.today()

        self.BASE_PATH = os.path.dirname(os.path.dirname(os.path.abspath(__file__).replace('/applications', '')))
        self.COMMON_PATH = os.path.join(self.BASE_PATH, 'common_docs/')
        self.MEDIA_PATH = os.path.join(self.BASE_PATH, 'media/')

        self.DEFAUL_DIRECTORY = f'{self.COMMON_PATH}documents/'
        self.OUTPUT_DIRECTORY = f'{self.MEDIA_PATH}documents/'

    def generate_filename(self):
        filename = f'{self.DEFAUL_DIRECTORY}cv.docx'
        output_filename = f'{self.OUTPUT_DIRECTORY}{self.profile.last_name}-{self.profile.first_name}-cv.docx'
        return filename, output_filename

    def get_lang_level(self, grade):
        level_grades = {
            '1': 'fließend',
            '2': 'sehr gut',
            '3': 'gut',
            '4': 'schlecht',
        }
        return level_grades.get(grade, None)

    def generate_document(self, fname, output_fname):
        document = Document(fname)

        from docx.shared import Cm

        from io import BytesIO

        output = BytesIO()

        if self.profile.photo_for_schengen is not None:
            try:
                response = requests.get(self.profile.photo_for_schengen.url)
                image_bytes = BytesIO(response.content)
                image = Image.open(image_bytes)
                new_size = (132, 170)
                image.resize(new_size)

                image.save(output, image.format)

                p = document.tables[0].rows[0].cells[0].add_paragraph()
                r = p.add_run()
                r.add_picture(output, width=Cm(3.5), height=Cm(4.5))
            except:
                pass

        style = document.styles['Normal']
        font = style.font
        font.name = 'Open Sans'
        font.size = Pt(11)

        my_styles = document.styles
        p_style = my_styles.add_style('custom-paragraph-style', WD_STYLE_TYPE.PARAGRAPH)
        p_style.base_style = my_styles['Normal']
        p_style.paragraph_format.space_before = Pt(2)
        p_style.paragraph_format.space_after = Pt(2)
        ch_style = my_styles.add_style('new-character-style', WD_STYLE_TYPE.CHARACTER)
        ch_style.base_style = my_styles['Default Paragraph Font']
        ch_style.font.name = 'Open Sans'
        ch_style.font.size = Pt(12)
        ch_style.font.bold = True

        full_name = document.paragraphs[0].text.replace('Mamatova Gulnaz', f'{self.profile.last_name} {self.profile.first_name}')
        document.paragraphs[0].text = full_name

        postcode_city = '720000 Bishkek' if self.profile.reg_city_en.lower() != 'osh' else '723500 Osh'
        post_code = document.paragraphs[1].text.replace('720000 Bischkek', postcode_city)
        document.paragraphs[1].text = post_code

        street = document.paragraphs[2].text.replace('Tynaliev 23/9', self.profile.live_street_number_translit)
        document.paragraphs[2].text = street

        phone = document.paragraphs[3].text.replace('+966707322502', self.profile.user.phone)
        document.paragraphs[3].text = phone

        email = document.paragraphs[4].text.replace('mamatovagulnaz64@gmail.com', self.profile.user.email)
        document.paragraphs[4].text = email

        birth_date = document.paragraphs[5].text.replace('27.12.1999', self.profile.bday.strftime('%d.%m.%Y'))
        document.paragraphs[5].text = birth_date

        new_paragraph = document.paragraphs[9].insert_paragraph_before(style='custom-paragraph-style')
        new_paragraph.add_run(f'{self.profile.start_date1.strftime("%d.%m.%Y")}-{self.profile.end_date1.strftime("%d.%m.%Y")}                {self.profile.company1}, {self.profile.country1}', style='new-character-style')

        new_paragraph = document.paragraphs[10].insert_paragraph_before(style='custom-paragraph-style')
        new_paragraph.add_run(f'                                                         {self.profile.position1}', style='new-character-style')


        new_paragraph = document.paragraphs[11].insert_paragraph_before(style='custom-paragraph-style')
        new_paragraph.add_run(f'{self.profile.start_date2.strftime("%d.%m.%Y")}-{self.profile.end_date2.strftime("%d.%m.%Y")}                {self.profile.company2}, {self.profile.country2}', style='new-character-style')

        new_paragraph = document.paragraphs[12].insert_paragraph_before(style='custom-paragraph-style')
        new_paragraph.add_run(f'                                                         {self.profile.position2}', style='new-character-style')


        row_nums = [16, 17, 18]

        if self.profile.company3 and self.profile.position3 and self.profile.country3 and self.profile.start_date3 and self.profile.end_date3:
            row_nums = [18, 19, 20]

            new_paragraph = document.paragraphs[13].insert_paragraph_before(style='custom-paragraph-style')
            new_paragraph.add_run(f'{self.profile.start_date3.strftime("%d.%m.%Y")}-{self.profile.end_date3.strftime("%d.%m.%Y")}                {self.profile.company3}, {self.profile.country3}', style='new-character-style')

            new_paragraph = document.paragraphs[14].insert_paragraph_before(style='custom-paragraph-style')
            new_paragraph.add_run(f'                                                         {self.profile.position3}', style='new-character-style')


        new_paragraph = document.paragraphs[row_nums[0]].insert_paragraph_before(style='custom-paragraph-style')
        new_paragraph.add_run('Bachelorstudium', style='new-character-style')

        new_paragraph = document.paragraphs[row_nums[1]].insert_paragraph_before(style='custom-paragraph-style')
        new_paragraph.add_run(f'{self.profile.study_start.strftime("%m/%Y")} - {self.profile.study_end.strftime("%m/%Y")}                       {self.profile.university.name_de}', style='new-character-style')

        new_paragraph = document.paragraphs[row_nums[2]].insert_paragraph_before(style='custom-paragraph-style')
        new_paragraph.add_run(f'                                                         {self.profile.faculty.name_de}', style='new-character-style')

        new_paragraph = document.add_paragraph(style='custom-paragraph-style')
        new_paragraph.add_run('Sprachkenntnisse', style='new-character-style')

        german_level = self.get_lang_level(self.profile.german)
        if german_level:
            new_paragraph = document.add_paragraph()
            new_paragraph.add_run(f'                                                           Deutsch - {german_level}')

        english_level = self.get_lang_level(self.profile.english)
        if english_level:
            new_paragraph = document.add_paragraph()
            new_paragraph.add_run(f'                                                           Englisch - {english_level}')

        turkish_level = self.get_lang_level(self.profile.turkish)
        if turkish_level:
            new_paragraph = document.add_paragraph()
            new_paragraph.add_run(f'                                                           Türkisch - {turkish_level}')

        russian_level = self.get_lang_level(self.profile.russian)
        if russian_level:
            new_paragraph = document.add_paragraph()
            new_paragraph.add_run(f'                                                           Russisch - {russian_level}')

        chinese_level = self.get_lang_level(self.profile.chinese)
        if chinese_level:
            new_paragraph = document.add_paragraph()
            new_paragraph.add_run(f'                                                           Chinesische - {chinese_level}')

        new_paragraph = document.add_paragraph(style='custom-paragraph-style')
        new_paragraph.add_run('EDV                                                   Microsoft Word, Excel, Power Point', style='new-character-style')

        new_paragraph = document.add_paragraph(style='custom-paragraph-style')

        driver_info = 'Nein' if not self.profile.driver_license else f'Kategorie {", ".join(self.profile.get_drive_categories_wo_tractor)}'
        new_paragraph.add_run(f'Führerschein                                    {driver_info}', style='new-character-style')

        new_paragraph = document.add_paragraph(style='custom-paragraph-style')
        new_paragraph.add_run(f'Hobby                                               {", ".join(self.profile.get_hobbies_de[:5])}', style='new-character-style')

        new_paragraph = document.add_paragraph()
        new_paragraph.add_run()

        new_paragraph = document.add_paragraph()
        new_paragraph.add_run(f'Bischkek, {self.today.strftime("%d.%m.%Y")}')
        document.save(output_fname)

    def generate_cv(self):
        if self.profile.is_confirmed and self.profile.is_admin_confirmed and self.profile.is_form_completed \
                    and self.profile.photo_for_schengen:
            filename, output_filename = self.generate_filename()
            self.generate_document(filename, output_filename)
            return output_filename
        return None


class AnketaGenerator:
    def __init__(self, profile):
        self.profile = profile

        filename_pdf = f'{self.profile.last_name}-{self.profile.first_name}-anketa.pdf'

        self.BASE_PATH = os.path.dirname(os.path.dirname(os.path.abspath(__file__).replace('/applications', '')))
        self.COMMON_PATH = os.path.join(self.BASE_PATH, 'common_docs/')
        self.MEDIA_PATH = os.path.join(self.BASE_PATH, 'media/')

        self.TEMPLATE_PATH = f'{self.COMMON_PATH}documents/anket.pdf'

        self.OUTPUT_PATH = f'{self.MEDIA_PATH}documents/anketa.pdf'

        self.FINAL_PDF_PATH = f'{self.MEDIA_PATH}documents/{filename_pdf}'

    def run(self, template_path, filename):
        canvas_data = self.get_overlay_canvas()
        form = self.merge(canvas_data, template_path=template_path)
        self.save(form, filename=filename)
        return filename

    def merge(self, overlay_canvas: io.BytesIO, template_path: str) -> io.BytesIO:
        template_pdf = pdfrw.PdfReader(template_path)
        overlay_pdf = pdfrw.PdfReader(overlay_canvas)
        for page, data in zip(template_pdf.pages, overlay_pdf.pages):
            overlay = pdfrw.PageMerge().add(data)[0]
            pdfrw.PageMerge(page).add(overlay).render()
        form = io.BytesIO()
        pdfrw.PdfWriter().write(form, template_pdf)
        form.seek(0)
        return form

    def get_overlay_canvas(self) -> io.BytesIO:

        data = io.BytesIO()
        pdf = canvas.Canvas(data)
        pdfmetrics.registerFont(TTFont('Roboto-Light', f'{self.COMMON_PATH}fonts/Roboto-Light.ttf'))
        pdfmetrics.registerFont(TTFont('Roboto-Regular', f'{self.COMMON_PATH}fonts/Roboto-Regular.ttf'))
        pdfmetrics.registerFont(TTFont('Roboto-Italic', f'{self.COMMON_PATH}fonts/Roboto-Italic.ttf'))
        pdfmetrics.registerFont(TTFont('Roboto-Thin', f'{self.COMMON_PATH}fonts/Roboto-Thin.ttf'))
        pdfmetrics.registerFont(TTFont('Roboto-Bold', f'{self.COMMON_PATH}fonts/Roboto-Bold.ttf'))

        pdf.setFont('Roboto-Regular', 12)

        avatar = ImageReader(self.profile.photo.url)
        pdf.drawImage(avatar, 375, 610, 120, 155)
        pdf.drawString(x=155, y=740, text=self.profile.last_name_ru)
        pdf.drawString(x=155, y=710, text=self.profile.first_name_ru)
        pdf.drawString(x=155, y=680, text=self.profile.bday.strftime('%d.%m.%Y'))
        pdf.drawString(x=155, y=650, text=self.profile.get_birth_country_display())
        pdf.drawString(x=155, y=620, text=self.profile.user.phone)
        pdf.drawString(x=155, y=588, text=self.profile.user.email)
        pdf.drawString(x=155, y=555, text=self.profile.live_street_number)
        pdf.drawString(x=155, y=522, text=self.profile.university.name_ru)
        pdf.drawString(x=155, y=487, text=self.profile.faculty.name_ru)
        pdf.drawString(x=155, y=454, text=str(self.profile.year) if self.profile.year else '')
        pdf.drawString(x=230, y=422, text=self.profile.get_shirt_size_display() if self.profile.shirt_size else '')
        pdf.drawString(x=340, y=422, text=str(self.profile.pants_size) if self.profile.pants_size else '')
        pdf.drawString(x=455, y=422, text=str(self.profile.shoe_size) if self.profile.shoe_size else '')

        driver_info = f'Есть, категории: {", ".join(self.profile.get_drive_categories_wo_tractor)}, стаж: {self.profile.get_driving_experience_display()}' if self.profile.driver_license else 'Нет'
        pdf.drawString(x=225, y=390, text=driver_info)

        pdf.drawString(x=225, y=358, text=self.profile.get_bicycle_skill_display())

        germany_experience = 'Да, был.' if self.profile.country1 == 'Deutschland' or self.profile.country2 == 'Deutschland' or self.profile.country3 == 'Deutschland' else 'Нет, не был.'
        pdf.drawString(x=225, y=325, text=germany_experience)
        pdf.drawString(x=180, y=291, text=self.profile.father_phone)
        pdf.drawString(x=415, y=291, text=self.profile.mother_phone)
        pdf.drawString(x=180, y=258, text=self.profile.father_company)
        pdf.drawString(x=415, y=258, text=self.profile.mother_company)
        pdf.drawString(x=150, y=223, text=self.profile.passport_number)
        pdf.drawString(x=150, y=206, text=self.profile.zagranpassport_number)
        pdf.drawString(x=140, y=158, text=f'{self.profile.start_date1.strftime("%d.%m.%Y")}-{self.profile.end_date1.strftime("%d.%m.%Y")}')
        pdf.drawString(x=273, y=158, text=self.profile.company1)
        pdf.drawString(x=379, y=158, text=self.profile.get_country1_display())
        pdf.drawString(x=494, y=158, text=self.profile.get_position1_display())
        pdf.drawString(x=140, y=126, text=f'{self.profile.start_date2.strftime("%d.%m.%Y")}-{self.profile.end_date2.strftime("%d.%m.%Y")}')
        pdf.drawString(x=273, y=126, text=self.profile.company2)
        pdf.drawString(x=379, y=126, text=self.profile.get_country2_display())
        pdf.drawString(x=494, y=126, text=self.profile.get_position2_display())

        if self.profile.company3 and self.profile.position3 and\
                self.profile.country3 and self.profile.start_date3 and self.profile.end_date3:
            pdf.drawString(x=140, y=93, text=f'{self.profile.start_date3.strftime("%d.%m.%Y")}-{self.profile.end_date3.strftime("%d.%m.%Y")}')
            pdf.drawString(x=273, y=93, text=self.profile.company3)
            pdf.drawString(x=379, y=93, text=self.profile.get_country3_display())
            pdf.drawString(x=494, y=93, text=self.profile.get_position3_display())

        pdf.save()
        data.seek(0)
        return data

    def save(self, form: io.BytesIO, filename: str):
        with open(filename, 'wb') as f:
            f.write(form.read())

    def generate_anketa(self):
        if self.profile.is_confirmed and self.profile.is_admin_confirmed and self.profile.is_form_completed\
                and self.profile.photo:
            template_path = self.TEMPLATE_PATH
            filename = self.FINAL_PDF_PATH
            return self.run(template_path, filename)
        return None
